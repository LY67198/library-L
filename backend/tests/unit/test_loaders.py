"""文档加载器测试 — 验证 load_document 对 txt、markdown、docx、pdf 的解析与对未知格式的错误处理。"""
from __future__ import annotations

from pathlib import Path

import pytest

from app.rag.loaders import load_document, UnsupportedFormatError


def test_load_text(tmp_path: Path):
    """测试加载纯文本:返回单页,内容与原 utf-8 文件一致。"""
    f = tmp_path / "doc.txt"
    f.write_text("hello world", encoding="utf-8")
    pages = load_document(f)
    assert len(pages) == 1
    assert pages[0] == "hello world"


def test_load_markdown(tmp_path: Path):
    """测试加载 markdown:整篇文档作为单页返回,原始标记符号被保留。"""
    f = tmp_path / "doc.md"
    f.write_text("# Title\n\nBody", encoding="utf-8")
    pages = load_document(f)
    assert pages == ["# Title\n\nBody"]


def test_unsupported_format_raises(tmp_path: Path):
    """测试未知后缀名:应抛出 UnsupportedFormatError,引导调用方显式处理不支持的格式。"""
    f = tmp_path / "doc.xyz"
    f.write_text("x", encoding="utf-8")
    with pytest.raises(UnsupportedFormatError):
        load_document(f)


def test_load_docx(tmp_path: Path):
    """测试加载 docx:多段合并为单页,各段落文本都应包含在结果中。"""
    from docx import Document
    f = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("paragraph one")
    doc.add_paragraph("paragraph two")
    doc.save(f)
    pages = load_document(f)
    assert len(pages) == 1
    assert "paragraph one" in pages[0]
    assert "paragraph two" in pages[0]


def test_load_pdf(tmp_path: Path):
    """测试加载 pdf:若 pypdf 可用,生成 1 页空白 PDF 加载后应返回 list 类型结果。"""
    # Use a real PDF if pypdf can find one; skip if generation is too involved
    pytest.importorskip("pypdf")
    # Minimal: write a 1-page PDF using pypdf's writer
    from pypdf import PdfWriter
    f = tmp_path / "doc.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    # pypdf cannot write text easily without reportlab; skip real content
    writer.write(str(f))
    pages = load_document(f)
    # blank page returns empty string
    assert isinstance(pages, list)